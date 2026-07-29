from __future__ import annotations

import json
import re
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIRS = {
    "fr": ROOT / "ressources" / "docx",
    "en": ROOT / "en" / "ressources" / "docx",
    "de": ROOT / "de" / "ressources" / "docx",
    "it": ROOT / "it" / "ressources" / "docx",
}
KINDS = ("fiche-eleve", "guide-encadrement", "supports", "corrige", "kit-complet")
FORBIDDEN = re.compile(r"(?:file:///|[A-Z]:\\|_private-inputs|(?:^|[/\\])tmp[/\\])", re.IGNORECASE)
QA_DIR = ROOT / "docs" / "qa" / "docx-imprimables-professionnels"


def has_child(element, tag: str) -> bool:
    properties = element.find(qn("w:pPr"))
    return properties is not None and properties.find(qn(tag)) is not None


def audit_document(path: Path, language: str) -> dict:
    problems: list[str] = []
    document = Document(path)
    paragraphs = list(document.paragraphs)
    text = "\n".join(paragraph.text for paragraph in paragraphs)
    for table in document.tables:
        for row in table.rows:
            text += "\n" + " | ".join(cell.text for cell in row.cells)

    if not text.strip():
        problems.append("document vide")
    if FORBIDDEN.search(text):
        problems.append("chemin privé ou local détecté")

    for section in document.sections:
        width_mm = section.page_width.mm
        height_mm = section.page_height.mm
        if abs(width_mm - 210) > 0.5 or abs(height_mm - 297) > 0.5:
            problems.append(f"format non A4: {width_mm:.1f} x {height_mm:.1f} mm")

    list_paragraphs = [
        paragraph
        for paragraph in paragraphs
        if paragraph.style and paragraph.style.name in ("List Bullet", "List Number")
    ]
    for paragraph in list_paragraphs:
        if not has_child(paragraph._p, "w:keepLines"):
            problems.append("élément de liste scindable")
            break

    rows = [row for table in document.tables for row in table.rows]
    for row in rows:
        row_properties = row._tr.find(qn("w:trPr"))
        if row_properties is None or row_properties.find(qn("w:cantSplit")) is None:
            problems.append("ligne de tableau scindable")
            break

    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        relationships = "".join(
            archive.read(name).decode("utf-8")
            for name in archive.namelist()
            if name.endswith(".rels")
        )
    if FORBIDDEN.search(document_xml) or FORBIDDEN.search(relationships):
        problems.append("chemin privé ou local détecté dans le paquet")

    return {
        "file": path.relative_to(ROOT).as_posix(),
        "language": language,
        "bytes": path.stat().st_size,
        "paragraphs": len(paragraphs),
        "tables": len(document.tables),
        "table_rows": len(rows),
        "list_items": len(list_paragraphs),
        "problems": problems,
    }


def main() -> None:
    results = []
    missing = []
    for language, directory in OUTPUT_DIRS.items():
        for activity_id in ("01", "02", "03", "04"):
            for kind in KINDS:
                path = directory / f"activite-{activity_id}-{kind}.docx"
                if not path.exists():
                    missing.append(path.relative_to(ROOT).as_posix())
                    continue
                results.append(audit_document(path, language))

    failures = [result for result in results if result["problems"]]
    summary = {
        "expected": 80,
        "documents": len(results),
        "missing": missing,
        "failures": failures,
        "by_language": {
            language: sum(result["language"] == language for result in results)
            for language in OUTPUT_DIRS
        },
    }
    QA_DIR.mkdir(parents=True, exist_ok=True)
    (QA_DIR / "test-results.json").write_text(
        json.dumps({"summary": summary, "documents": results}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    print(json.dumps(summary, ensure_ascii=False))
    if missing or failures:
        raise SystemExit(1)


if __name__ == "__main__":
    main()
