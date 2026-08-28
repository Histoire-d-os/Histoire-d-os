from __future__ import annotations

import argparse
import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from lxml import html
from PIL import Image


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "ressources" / "print" / "pdf-sources"
ACCENTS = {"01": "9A5B25", "02": "356F62", "03": "87444C", "04": "315F79"}
OUTPUTS = {
    "fiche": "fiche-eleve",
    "guide": "guide-encadrement",
    "supports": "supports",
    "corrige": "corrige",
    "kit": "kit-complet",
}
OUTPUT_DIRS = {
    "fr": ROOT / "ressources" / "docx",
    "en": ROOT / "en" / "ressources" / "docx",
    "de": ROOT / "de" / "ressources" / "docx",
    "it": ROOT / "it" / "ressources" / "docx",
}
LANGUAGE_TAGS = {"fr": "fr-FR", "en": "en-US", "de": "de-DE", "it": "it-IT"}
LABELS = {
    "fr": {"pilot": "Version pilote", "trace": "Traçabilité", "trace_link": "Page de l'activité, mises à jour et crédits en ligne", "page": "Page"},
    "en": {"pilot": "Pilot version", "trace": "Traceability", "trace_link": "Activity page, updates and online credits", "page": "Page"},
    "de": {"pilot": "Pilotversion", "trace": "Nachvollziehbarkeit", "trace_link": "Aktivitätsseite, Aktualisierungen und Online-Bildnachweise", "page": "Seite"},
    "it": {"pilot": "Versione pilota", "trace": "Tracciabilità", "trace_link": "Pagina dell'attività, aggiornamenti e crediti online", "page": "Pagina"},
}
METADATA = {
    "fr": {
        "activity": "Activité",
        "subject": "Ressource pédagogique pilote sur le talus/astragale",
    },
    "en": {
        "activity": "Activity",
        "subject": "Pilot educational resource about the talus/astragalus",
    },
    "de": {
        "activity": "Aktivität",
        "subject": "Pädagogische Pilotressource zum Talus/Astragalus",
    },
    "it": {
        "activity": "Attività",
        "subject": "Risorsa educativa pilota sul talo/astragalo",
    },
}
KIND_LABELS = {
    "fr": {
        "fiche": "fiche élève",
        "guide": "guide d'encadrement",
        "supports": "supports",
        "corrige": "corrigé",
        "kit": "kit complet",
    },
    "en": {
        "fiche": "learner sheet",
        "guide": "facilitator guide",
        "supports": "supports",
        "corrige": "answer guidance",
        "kit": "complete kit",
    },
    "de": {
        "fiche": "Arbeitsblatt",
        "guide": "Leitfaden",
        "supports": "Materialien",
        "corrige": "Lösungshinweise",
        "kit": "Komplettpaket",
    },
    "it": {
        "fiche": "scheda allievo",
        "guide": "guida per la conduzione",
        "supports": "materiali",
        "corrige": "indicazioni per la correzione",
        "kit": "kit completo",
    },
}
ACTIVITY_SLUGS = {
    "01": "activite-01-trouver-astragale.html",
    "02": "activite-02-associer-astragale-animal.html",
    "03": "activite-03-achille-histoire-transformee.html",
    "04": "activite-04-un-os-plusieurs-vies.html",
}
CONTENT_WIDTH_DXA = 10318


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_fixed_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        existing = tbl_pr.find(qn(tag))
        if existing is not None:
            tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_paragraph_fill(paragraph, fill: str, border: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border:
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border)
        p_bdr.append(left)
        p_pr.append(p_bdr)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text_node = OxmlElement("w:t")
    text_node.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instr, separate, text_node, end))


def set_document_language(document: Document, language_tag: str) -> None:
    styles = document.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)
    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is None:
        r_pr_default = OxmlElement("w:rPrDefault")
        doc_defaults.append(r_pr_default)
    r_pr = r_pr_default.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r_pr_default.append(r_pr)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    for attribute in ("w:val", "w:eastAsia", "w:bidi"):
        lang.set(qn(attribute), language_tag)

    for style in document.styles:
        style_r_pr = style.element.get_or_add_rPr()
        style_lang = style_r_pr.find(qn("w:lang"))
        if style_lang is None:
            style_lang = OxmlElement("w:lang")
            style_r_pr.append(style_lang)
        for attribute in ("w:val", "w:eastAsia", "w:bidi"):
            style_lang.set(qn(attribute), language_tag)


def add_hyperlink(paragraph, text: str, url: str, color: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    run_color = OxmlElement("w:color")
    run_color.set(qn("w:val"), color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend((run_color, underline))
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend((r_pr, text_node))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_numbering(document: Document, ordered: bool) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    level.extend((start, num_fmt, lvl_text))
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend((tabs, ind, spacing))
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = numbering.add_num(abstract_id)
    return int(num.numId)


class DocxBuilder:
    def __init__(self, language: str, activity_id: str, kind: str, source: Path):
        self.language = language
        self.activity_id = activity_id
        self.kind = kind
        self.source = source
        self.accent = ACCENTS[activity_id]
        self.in_cover = False
        self.document = Document()
        self.configure_document()

    def configure_document(self) -> None:
        document = self.document
        language_tag = LANGUAGE_TAGS[self.language]
        set_document_language(document, language_tag)
        section = document.sections[0]
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(15)
        section.right_margin = Mm(14)
        section.bottom_margin = Mm(18)
        section.left_margin = Mm(14)
        section.header_distance = Mm(7)
        section.footer_distance = Mm(8)

        normal = document.styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = RGBColor.from_string("1F2523")
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.18

        for name, size, before, after, color in (
            ("Title", 28, 0, 12, self.accent),
            ("Heading 1", 22, 0, 10, "1F2523"),
            ("Heading 2", 14, 12, 6, self.accent),
            ("Heading 3", 11.5, 8, 4, "1F2523"),
        ):
            style = document.styles[name]
            style.font.name = "Georgia"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor.from_string(color)
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = True

        for name in ("Caption",):
            style = document.styles[name]
            style.font.name = "Arial"
            style.font.size = Pt(8.5)
            style.font.color.rgb = RGBColor.from_string("56605C")
            style.paragraph_format.space_after = Pt(5)
            style.paragraph_format.keep_together = True

        for name in ("List Bullet", "List Number"):
            style = document.styles[name]
            style.font.name = "Arial"
            style.font.size = Pt(10.5)
            style.paragraph_format.left_indent = Mm(9.5)
            style.paragraph_format.first_line_indent = Mm(-4.75)
            style.paragraph_format.space_after = Pt(4)
            style.paragraph_format.line_spacing = 1.18

        if "Resource Credit" not in document.styles:
            style = document.styles.add_style("Resource Credit", WD_STYLE_TYPE.PARAGRAPH)
        style = document.styles["Resource Credit"]
        style.font.name = "Arial"
        style.font.size = Pt(8)
        style.font.color.rgb = RGBColor.from_string("56605C")
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.keep_together = True

        core = document.core_properties
        core.author = "Histoires d'os"
        metadata = METADATA[self.language]
        core.title = (
            f"{metadata['activity']} {self.activity_id} - "
            f"{KIND_LABELS[self.language][self.kind]} ({self.language.upper()})"
        )
        core.subject = metadata["subject"]
        core.keywords = (
            f"Histoires d'os, talus, astragale, {self.language}, "
            f"{metadata['activity']} {self.activity_id}"
        )

        header = section.header.paragraphs[0]
        header.text = "Histoires d'os"
        header.style = document.styles["Resource Credit"]
        header.runs[0].bold = True
        header.paragraph_format.space_after = Pt(0)

        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.style = document.styles["Resource Credit"]
        footer.add_run(f"{LABELS[self.language]['pilot']} · {LABELS[self.language]['page']} ")
        add_field(footer, "PAGE")
        footer.add_run(" / ")
        add_field(footer, "NUMPAGES")

        settings = document.settings.element
        theme = settings.find(qn("w:themeFontLang"))
        if theme is None:
            theme = OxmlElement("w:themeFontLang")
            settings.append(theme)
        theme.set(qn("w:val"), language_tag)
        theme.set(qn("w:eastAsia"), language_tag)
        theme.set(qn("w:bidi"), language_tag)

    def add_inline(self, paragraph, node, bold=False, italic=False) -> None:
        def append_text(value: str | None, is_bold: bool, is_italic: bool) -> None:
            if value:
                run = paragraph.add_run(value)
                run.bold = is_bold
                run.italic = is_italic

        append_text(node.text, bold or node.tag in ("strong", "b"), italic or node.tag in ("em", "i"))
        for child in node:
            child_bold = bold or child.tag in ("strong", "b")
            child_italic = italic or child.tag in ("em", "i")
            if child.tag == "br":
                paragraph.add_run().add_break()
            elif child.tag in ("p", "div"):
                if paragraph.text:
                    paragraph.add_run().add_break()
                self.add_inline(paragraph, child, child_bold, child_italic)
            elif child.tag == "a" and child.get("href"):
                add_hyperlink(paragraph, " ".join(child.text_content().split()), child.get("href"), self.accent)
            else:
                self.add_inline(paragraph, child, child_bold, child_italic)
            append_text(child.tail, bold, italic)

    def add_paragraph(self, node, style=None, fill=None) -> object:
        paragraph = self.document.add_paragraph(style=style)
        self.add_inline(paragraph, node)
        paragraph.paragraph_format.keep_together = True
        if fill:
            set_paragraph_fill(paragraph, fill, self.accent)
            paragraph.paragraph_format.left_indent = Mm(3)
            paragraph.paragraph_format.right_indent = Mm(2)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(6)
        return paragraph

    def add_heading(self, node) -> object:
        level = {"h1": 1, "h2": 2, "h3": 3}[node.tag]
        cover_title = self.in_cover and node.tag == "h1"
        paragraph = self.document.add_paragraph(style=f"Heading {level}")
        self.add_inline(paragraph, node)
        if cover_title:
            paragraph.paragraph_format.space_after = Pt(12)
            for run in paragraph.runs:
                run.font.name = "Georgia"
                run.font.size = Pt(28)
                run.font.bold = True
                run.font.color.rgb = RGBColor.from_string(self.accent)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
        classes = set((node.get("class") or "").split())
        if "page-break-before" in classes:
            paragraph.paragraph_format.page_break_before = True
        return paragraph

    def add_list(self, node) -> None:
        items = node.xpath("./li")
        style = "List Number" if node.tag == "ol" else "List Bullet"
        for index, item in enumerate(items):
            paragraph = self.document.add_paragraph(style=style)
            self.add_inline(paragraph, item)
            paragraph.paragraph_format.keep_together = True
            paragraph.paragraph_format.keep_with_next = index < len(items) - 1

    def column_widths(self, count: int, rows) -> list[int]:
        if count == 2:
            ratios = [0.30, 0.70]
        elif count == 3:
            ratios = [0.24, 0.38, 0.38]
        elif count == 4:
            ratios = [0.18, 0.28, 0.28, 0.26]
        else:
            ratios = [1 / count] * count
        widths = [round(CONTENT_WIDTH_DXA * ratio) for ratio in ratios]
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
        return widths

    def add_table(self, node) -> None:
        html_rows = node.xpath("./thead/tr | ./tbody/tr | ./tr")
        if not html_rows:
            return
        count = max(len(row.xpath("./th | ./td")) for row in html_rows)
        table = self.document.add_table(rows=len(html_rows), cols=count)
        table.style = "Table Grid"
        widths = self.column_widths(count, html_rows)
        for row_index, html_row in enumerate(html_rows):
            word_row = table.rows[row_index]
            prevent_row_split(word_row)
            cells = html_row.xpath("./th | ./td")
            if row_index == 0 and (html_row.getparent().tag == "thead" or all(cell.tag == "th" for cell in cells)):
                set_repeat_table_header(word_row)
            for column, html_cell in enumerate(cells):
                cell = word_row.cells[column]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph = cell.paragraphs[0]
                self.add_inline(paragraph, html_cell)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                paragraph.paragraph_format.keep_together = True
                if html_cell.tag == "th" or row_index == 0:
                    set_cell_shading(cell, "F0F1EE")
                    for run in paragraph.runs:
                        run.bold = True
                style = html_cell.get("style") or ""
                match = re.search(r"height\s*:\s*(\d+)mm", style)
                if match:
                    word_row.height = Mm(int(match.group(1)))
                    word_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        set_fixed_table_geometry(table, widths)
        if len(html_rows) <= 6 and len(" ".join(node.text_content().split())) < 1800:
            paragraphs = [p for row in table.rows for cell in row.cells for p in cell.paragraphs]
            for paragraph in paragraphs[:-1]:
                paragraph.paragraph_format.keep_with_next = True
        after = self.document.add_paragraph()
        after.paragraph_format.space_after = Pt(0)

    def add_image(self, node, max_width_mm=158, container=None) -> None:
        src = node.get("src")
        if not src or src.startswith(("http://", "https://", "data:")):
            return
        image_path = (self.source.parent / src).resolve()
        if not image_path.exists() or not image_path.is_file():
            return
        with Image.open(image_path) as image:
            width, height = image.size
        max_height_mm = 112
        width_mm = min(max_width_mm, max_height_mm * width / max(height, 1))
        paragraph = container.add_paragraph() if container is not None else self.document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_together = True
        inline_shape = paragraph.add_run().add_picture(str(image_path), width=Mm(width_mm))
        alt_text = " ".join((node.get("alt") or "Illustration pédagogique").split())
        inline_shape._inline.docPr.set("descr", alt_text)
        inline_shape._inline.docPr.set("title", alt_text[:120])

    def add_answer_zone(self, node) -> None:
        line_count = 3 if "compact" in (node.get("class") or "") else 5
        for _ in range(line_count):
            paragraph = self.document.add_paragraph(" ")
            paragraph.paragraph_format.space_after = Pt(5)
            paragraph.paragraph_format.keep_together = True
            p_pr = paragraph._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "dotted")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "2")
            bottom.set(qn("w:color"), "A9B0AC")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)

    def add_grid(self, node) -> None:
        cards = node.xpath("./article")
        columns = 3 if "quick-evaluation" in (node.get("class") or "") else 2
        groups = (len(cards) + columns - 1) // columns
        table = self.document.add_table(rows=groups * 2, cols=columns)
        table.style = "Table Grid"
        widths = [CONTENT_WIDTH_DXA // columns] * columns
        widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
        for group in range(groups):
            header_row = table.rows[group * 2]
            content_row = table.rows[group * 2 + 1]
            prevent_row_split(header_row)
            prevent_row_split(content_row)
            set_repeat_table_header(header_row)
            for column in range(columns):
                index = group * columns + column
                header_cell = header_row.cells[column]
                content_cell = content_row.cells[column]
                for cell in (header_cell, content_cell):
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                    set_cell_margins(cell, top=130, bottom=130, start=150, end=150)
                set_cell_shading(header_cell, "F0F1EE")
                set_cell_shading(content_cell, "F7F4EF")
                if index >= len(cards):
                    continue
                card = cards[index]
                headings = card.xpath("./h1 | ./h2 | ./h3")
                if headings:
                    header_cell.paragraphs[0]._element.getparent().remove(header_cell.paragraphs[0]._element)
                    self.add_to_container_paragraph(header_cell, headings[0], "Heading 3")
                content_cell.paragraphs[0]._element.getparent().remove(content_cell.paragraphs[0]._element)
                body = copy.deepcopy(card)
                for heading in body.xpath("./h1 | ./h2 | ./h3"):
                    heading.getparent().remove(heading)
                self.add_children(body, container=content_cell, image_width=74 if columns == 2 else 48)
        set_fixed_table_geometry(table, widths)
        self.document.add_paragraph().paragraph_format.space_after = Pt(0)

    def add_to_container_paragraph(self, container, node, style=None):
        paragraph = container.add_paragraph(style=style)
        self.add_inline(paragraph, node)
        paragraph.paragraph_format.keep_together = True
        return paragraph

    def add_children(self, node, container=None, image_width=150) -> None:
        for child in node:
            classes = set((child.get("class") or "").split())
            if child.tag in ("h1", "h2", "h3"):
                if container is None:
                    self.add_heading(child)
                else:
                    style = {"h1": "Heading 1", "h2": "Heading 2", "h3": "Heading 3"}[child.tag]
                    self.add_to_container_paragraph(container, child, style)
            elif child.tag == "p":
                style = "Resource Credit" if "credit-line" in classes or "pdf-source" in classes else None
                if container is None:
                    self.add_paragraph(child, style=style)
                else:
                    self.add_to_container_paragraph(container, child, style)
            elif child.tag in ("ul", "ol"):
                if container is None:
                    self.add_list(child)
                else:
                    items = child.xpath("./li")
                    style = "List Number" if child.tag == "ol" else "List Bullet"
                    for index, item in enumerate(items):
                        paragraph = container.add_paragraph(style=style)
                        self.add_inline(paragraph, item)
                        paragraph.paragraph_format.keep_together = True
                        paragraph.paragraph_format.keep_with_next = index < len(items) - 1
            elif child.tag == "table":
                if container is None:
                    self.add_table(child)
                else:
                    paragraph = container.add_paragraph(" ".join(child.text_content().split()))
                    paragraph.paragraph_format.keep_together = True
            elif child.tag == "img":
                self.add_image(child, image_width, container=container)
            elif child.tag == "figure":
                image = child.xpath(".//img")
                if image:
                    self.add_image(image[0], image_width, container=container)
                caption = child.xpath("./figcaption")
                if caption:
                    if container is None:
                        self.add_paragraph(caption[0], style="Caption")
                    else:
                        self.add_to_container_paragraph(container, caption[0], "Caption")
            elif "answer-zone" in classes:
                self.add_answer_zone(child)
            elif "neutral-placeholder" in classes:
                paragraph = self.document.add_paragraph(" ".join(child.text_content().split()))
                paragraph.paragraph_format.keep_together = True
                set_paragraph_fill(paragraph, "F1F2F0", self.accent)
            elif classes.intersection({"visual-grid", "cut-grid", "pdf-grid-2", "quick-evaluation", "observation-hypothesis"}) and child.xpath("./article"):
                self.add_grid(child)
            elif child.tag == "dl" or "pdf-identification" in classes:
                pairs = child.xpath(".//div[dt and dd]")
                if pairs:
                    table = self.document.add_table(rows=2, cols=len(pairs))
                    widths = [CONTENT_WIDTH_DXA // len(pairs)] * len(pairs)
                    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
                    for row in table.rows:
                        prevent_row_split(row)
                    set_repeat_table_header(table.rows[0])
                    for index, pair in enumerate(pairs):
                        header_cell = table.cell(0, index)
                        value_cell = table.cell(1, index)
                        set_cell_shading(header_cell, "F0F1EE")
                        set_cell_shading(value_cell, "F7F4EF")
                        header_cell.paragraphs[0]._element.getparent().remove(header_cell.paragraphs[0]._element)
                        value_cell.paragraphs[0]._element.getparent().remove(value_cell.paragraphs[0]._element)
                        terms = pair.xpath("./dt")
                        values = pair.xpath("./dd")
                        if terms:
                            paragraph = header_cell.add_paragraph(style="Resource Credit")
                            self.add_inline(paragraph, terms[0])
                            for run in paragraph.runs:
                                run.bold = True
                        if values:
                            paragraph = value_cell.add_paragraph()
                            self.add_inline(paragraph, values[0])
                    set_fixed_table_geometry(table, widths)
            elif child.tag in ("div", "article", "section", "header", "nav", "aside"):
                if "mission-box" in classes or "pdf-callout" in classes or "word-bank" in classes or "pilot-status" in classes:
                    paragraph = self.document.add_paragraph()
                    self.add_inline(paragraph, child)
                    paragraph.paragraph_format.keep_together = True
                    set_paragraph_fill(paragraph, "F7F4EF", self.accent)
                else:
                    self.add_children(child, container=container, image_width=image_width)
            elif child.tag == "br":
                if container is None:
                    self.document.add_paragraph()

    def add_traceability(self) -> None:
        prefix = "" if self.language == "fr" else f"/{self.language}"
        url = f"https://histoire-d-os.github.io/Histoire-d-os{prefix}/ressources/{ACTIVITY_SLUGS[self.activity_id]}"
        paragraph = self.document.add_paragraph(style="Resource Credit")
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.add_run(f"{LABELS[self.language]['trace']} : ").bold = True
        add_hyperlink(paragraph, LABELS[self.language]["trace_link"], url, self.accent)

    def build(self, document_node, include_traceability=True) -> None:
        self.in_cover = document_node.get("data-document") == "cover"
        self.add_children(document_node)
        self.in_cover = False
        if include_traceability:
            self.add_traceability()

    def save(self, destination: Path) -> None:
        destination.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(destination)


def source_path(activity_id: str, language: str) -> Path:
    suffix = "" if language == "fr" else f".{language}"
    return SOURCE_DIR / f"activite-{activity_id}{suffix}.html"


def build_one(activity_id: str, language: str, output_kind: str) -> Path:
    source = source_path(activity_id, language)
    tree = html.fromstring(source.read_text(encoding="utf-8"))
    builder = DocxBuilder(language, activity_id, output_kind, source)
    if output_kind == "kit":
        order = ("cover", "guide", "fiche", "supports", "corrige", "sources")
        for index, kind in enumerate(order):
            node = tree.xpath(f'//section[@data-document="{kind}"]')[0]
            if index:
                builder.document.add_page_break()
            builder.build(node, include_traceability=kind == "sources")
    else:
        node = tree.xpath(f'//section[@data-document="{output_kind}"]')[0]
        builder.build(node)
    destination = OUTPUT_DIRS[language] / f"activite-{activity_id}-{OUTPUTS[output_kind]}.docx"
    builder.save(destination)
    return destination


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--languages", nargs="+", choices=sorted(OUTPUT_DIRS), default=sorted(OUTPUT_DIRS))
    parser.add_argument("--kinds", nargs="+", choices=sorted(OUTPUTS), default=list(OUTPUTS))
    parser.add_argument("--activities", nargs="+", choices=("01", "02", "03", "04"), default=("01", "02", "03", "04"))
    args = parser.parse_args()
    outputs = []
    for language in args.languages:
        for activity_id in args.activities:
            for kind in args.kinds:
                destination = build_one(activity_id, language, kind)
                outputs.append(destination)
                print(destination.relative_to(ROOT).as_posix())
    print(f"{len(outputs)} documents Word générés.")


if __name__ == "__main__":
    main()
